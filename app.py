import docx

# Step 1: Load Offer Letter DOCX Template
def load_template(docx_template_path):
    return docx.Document(docx_template_path)

# Step 2: Replace Placeholders in Paragraphs and Runs
def replace_placeholders(doc, hr_data):
    for para in doc.paragraphs:
        full_text = ''.join(run.text for run in para.runs)
        # Ensure the placeholders are properly formatted
        for key, value in hr_data.items():
            if key in full_text:
                full_text = full_text.replace(key, value)
        
        # Clear all runs and insert the modified full_text
        if len(para.runs) > 0:
            for run in para.runs:
                run.text = ""  # Clear the existing text in each run
            para.runs[0].text = full_text  # Insert the full text with replacements
    
    return doc

def add_benefits_section(doc, benefit_a, benefit_b, benefit_c):
    # Go through each paragraph and replace placeholders individually
    for para in doc.paragraphs:
        if "{benefit_A}" in para.text:
            para.text = para.text.replace("{benefit_A}", f"{benefit_a}")
        if "{benefit_B}" in para.text:
            para.text = para.text.replace("{benefit_B}", f"{benefit_b}")
        if "{benefit_C}" in para.text:
            para.text = para.text.replace("{benefit_C}", f"{benefit_c}")
    
    return doc




# Step 4: Save the customized DOCX
def save_customized_docx(doc, output_path):
    doc.save(output_path)
